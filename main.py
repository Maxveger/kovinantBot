from flask import Flask, request, send_file, redirect, url_for
import os
import shutil
import pandas as pd
from docx import Document
import zipfile
from datetime import datetime

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
ACTS_FOLDER = 'generated_acts'
CONTRACTS_FOLDER = 'generated_contracts'
TEMPLATES_FOLDER = 'templates'
TMP_FOLDER = 'tmp'
ACT_TEMPLATE = os.path.join(TEMPLATES_FOLDER, 'templ_akt.docx')
CONTRACT_TEMPLATE = os.path.join(TEMPLATES_FOLDER, 'templ_dogovor.docx')
ZIP_NAME = os.path.join(TMP_FOLDER, 'generated_documents.zip')

for folder in [UPLOAD_FOLDER, ACTS_FOLDER, CONTRACTS_FOLDER, TEMPLATES_FOLDER, TMP_FOLDER]:
    os.makedirs(folder, exist_ok=True)

ALLOWED_EXTENSIONS = {'xls', 'xlsx', 'docx'}
success_message = ""

def allowed_file(filename, types):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in types

@app.route('/')
def index():
    global success_message
    msg = f"<p style='color:green'>{success_message}</p>" if success_message else ''
    success_message = ''
    return f'''
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset='UTF-8'>
            <title>Загрузка Excel-файла для генерации документов</title>
        </head>
        <body>
            <h1>Загрузка Excel-файла для генерации документов</h1>
            <form method='POST' enctype='multipart/form-data' action='/upload'>
                <input type='file' name='file'><input type='submit' value='Отправить'>
            </form>
            <h2>Как пользоваться:</h2>
            <ul>
                <li>Подготовьте Excel-файл в формате .xls или .xlsx.</li>
                <li>В первой строке должны быть заголовки колонок (например: Имя, Дата, Адрес).</li>
                <li>Каждая строка ниже — это отдельная запись для генерации документов.</li>
                <li>Выберите файл и нажмите «Отправить».</li>
                <li>Через несколько секунд начнётся скачивание ZIP-архива с готовыми файлами.</li>
            </ul>
            <form method='POST' enctype='multipart/form-data' action='/upload_template'>
                <h3>Обновить шаблон акта:</h3>
                <input type='file' name='template_akt'><input type='submit' name='submit_type' value='Загрузить акт'><br><br>
                <h3>Обновить шаблон договора:</h3>
                <input type='file' name='template_dogovor'><input type='submit' name='submit_type' value='Загрузить договор'>
            </form>
            <p><b>Важно:</b> Убедитесь, что файл заполнен корректно, чтобы избежать ошибок при генерации.</p>
            {msg}
        </body>
        </html>
    '''

@app.route('/upload_template', methods=['POST'])
def upload_template():
    global success_message
    if 'submit_type' in request.form:
        if request.form['submit_type'] == 'Загрузить акт' and 'template_akt' in request.files:
            file = request.files['template_akt']
            if allowed_file(file.filename, {'docx'}):
                file.save(ACT_TEMPLATE)
                success_message = 'Шаблон акта обновлён.'
        elif request.form['submit_type'] == 'Загрузить договор' and 'template_dogovor' in request.files:
            file = request.files['template_dogovor']
            if allowed_file(file.filename, {'docx'}):
                file.save(CONTRACT_TEMPLATE)
                success_message = 'Шаблон договора обновлён.'
    return redirect(url_for('index'))

@app.route('/upload', methods=['POST'])
def upload_file():
    global success_message
    if 'file' not in request.files:
        success_message = 'Нет файла в запросе.'
        return redirect(url_for('index'))

    file = request.files['file']
    if file.filename == '':
        success_message = 'Файл не выбран.'
        return redirect(url_for('index'))

    if file and allowed_file(file.filename, {'xls', 'xlsx'}):
        filename = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(filename)
        return process_file(filename)
    else:
        success_message = 'Неверный формат файла.'
        return redirect(url_for('index'))

def replace_placeholders(doc: Document, data_row):
    for para in doc.paragraphs:
        for column in data_row.index:
            placeholder = f"{{{column}}}"
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, str(data_row[column]))

    for table in doc.tables:
        for row_cells in table.rows:
            for cell in row_cells.cells:
                for column in data_row.index:
                    placeholder = f"{{{column}}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(data_row[column]))

def process_file(file_path):
    global success_message

    data = pd.read_excel(file_path, dtype=str)
    if "date_pass" in data.columns:
        data["date_pass"] = pd.to_datetime(data["date_pass"], errors='coerce').dt.strftime("%d.%m.%Y")

    for index, row in data.iterrows():
        act_filename = os.path.join(ACTS_FOLDER, f"{row['name']}_act.docx")
        shutil.copy(ACT_TEMPLATE, act_filename)
        new_act = Document(act_filename)
        replace_placeholders(new_act, row)
        new_act.save(act_filename)

        contract_filename = os.path.join(CONTRACTS_FOLDER, f"{row['name']}_contract.docx")
        shutil.copy(CONTRACT_TEMPLATE, contract_filename)
        new_contract = Document(contract_filename)
        replace_placeholders(new_contract, row)
        new_contract.save(contract_filename)

    with zipfile.ZipFile(ZIP_NAME, 'w') as zipf:
        for folder in [ACTS_FOLDER, CONTRACTS_FOLDER]:
            for root, dirs, files in os.walk(folder):
                for file in files:
                    path = os.path.join(root, file)
                    zipf.write(path, os.path.relpath(path, folder))

    shutil.rmtree(ACTS_FOLDER)
    shutil.rmtree(CONTRACTS_FOLDER)
    os.makedirs(ACTS_FOLDER)
    os.makedirs(CONTRACTS_FOLDER)

    success_message = 'Документы сгенерированы. Архив загружается...'
    return send_file(ZIP_NAME, as_attachment=True)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.getenv('PORT', 5000)))
